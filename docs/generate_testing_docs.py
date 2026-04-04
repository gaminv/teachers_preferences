from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
TEMPLATE_DIR = Path(r"C:\Users\gamin\Downloads\Telega Desktop")
OUTPUT_DIR = ROOT / "docs"


def find_template(index: int) -> Path:
    return sorted([p for p in TEMPLATE_DIR.iterdir() if p.suffix == ".docx"])[index]


UNIT_TEMPLATE = TEMPLATE_DIR / "Unit-tests.docx"
INTEGRATION_TEMPLATE = find_template(5)
UNIT_OUTPUT = OUTPUT_DIR / "Модульное тестирование Teachers Preferences.docx"
INTEGRATION_OUTPUT = OUTPUT_DIR / "Интеграционное и системное тестирование Teachers Preferences.docx"


def reset_document(doc: Document) -> None:
    body = doc._element.body
    for element in list(body):
        if element.tag.endswith("sectPr"):
            continue
        body.remove(element)


def set_base_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(14)
    for style_name in ["Title", "Heading 1", "Heading 2", "Heading 3"]:
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"


def add_paragraph(doc: Document, text: str = "", *, bold: bool = False, align=None, style=None):
    paragraph = doc.add_paragraph(style=style)
    if text:
        run = paragraph.add_run(text)
        run.bold = bold
        run.font.name = "Times New Roman"
        run.font.size = Pt(14)
    if align is not None:
        paragraph.alignment = align
    paragraph.paragraph_format.first_line_indent = Pt(35)
    paragraph.paragraph_format.space_after = Pt(6)
    return paragraph


def add_bullet(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Pt(18)
    run = paragraph.add_run(f"• {text}")
    run.font.name = "Times New Roman"
    run.font.size = Pt(14)


_number_counter = 0


def add_numbered(doc: Document, text: str) -> None:
    global _number_counter
    _number_counter += 1
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Pt(18)
    run = paragraph.add_run(f"{_number_counter}. {text}")
    run.font.name = "Times New Roman"
    run.font.size = Pt(14)


def add_heading(doc: Document, text: str, level: int) -> None:
    heading = doc.add_paragraph(style=f"Heading {level}")
    run = heading.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(16 if level == 1 else 14)
    run.bold = True


def add_table(doc: Document, headers, rows) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, value in enumerate(headers):
        hdr[i].text = value
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(12)


def add_page_break(doc: Document) -> None:
    doc.add_page_break()


def build_title_page(doc: Document, subtitle: str) -> None:
    for line in [
        "Министерство науки и высшего образования РФ",
        "Санкт-Петербургский политехнический университет Петра Великого",
        "Институт компьютерных наук и кибербезопасности",
        "Высшая школа программной инженерии",
    ]:
        add_paragraph(doc, line, align=WD_ALIGN_PARAGRAPH.CENTER)

    for _ in range(6):
        add_paragraph(doc)

    add_paragraph(doc, "ОТЧЁТ ПО КУРСОВОМУ ПРОЕКТУ", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(doc, "по дисциплине «Технологии разработки качественного программного обеспечения»", align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(doc, subtitle, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

    for _ in range(4):
        add_paragraph(doc)

    add_paragraph(doc, "Выполнил: студент гр. 5130904/20103 Гамин В.В.", align=WD_ALIGN_PARAGRAPH.LEFT)
    add_paragraph(doc, "Руководитель: Маслаков А.П.", align=WD_ALIGN_PARAGRAPH.LEFT)

    for _ in range(7):
        add_paragraph(doc)

    add_paragraph(doc, "Санкт-Петербург", align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(doc, "2026 г.", align=WD_ALIGN_PARAGRAPH.CENTER)
    add_page_break(doc)


def build_unit_doc() -> None:
    global _number_counter
    doc = Document(UNIT_TEMPLATE)
    reset_document(doc)
    set_base_styles(doc)
    build_title_page(doc, "Модульное тестирование приложения «Teachers Preferences»")

    add_heading(doc, "1. Описание выполненной работы", 1)
    add_heading(doc, "1.1. Общая характеристика проекта", 2)
    add_paragraph(
        doc,
        "Приложение Teachers Preferences предназначено для сбора, хранения и анализа предпочтений преподавателей "
        "по учебной нагрузке. Архитектура решения включает web-интерфейс для преподавателя и администратора, "
        "backend API на Spring Boot, а также отдельный модуль schedule-etl для импорта и экспорта данных расписания."
    )
    add_table(
        doc,
        ["Компонент", "Технологии", "Назначение"],
        [
            ["backend", "Java 17 / Spring Boot / Maven", "REST API, аутентификация, работа с предпочтениями и экспортом Excel"],
            ["frontend", "React / Vite / Vitest", "Клиентский интерфейс преподавателя и администратора"],
            ["schedule-etl", "Java 17 / Maven / Apache POI", "Импорт, преобразование и выгрузка табличных данных"],
            ["db", "PostgreSQL", "Хранение учетных записей и предпочтений преподавателей"],
        ],
    )

    add_heading(doc, "1.2. Используемые инструменты", 2)
    add_table(
        doc,
        ["Инструмент", "Назначение"],
        [
            ["JUnit 5", "Модульные тесты backend и schedule-etl"],
            ["Mockito", "Mock-объекты зависимостей и изоляция unit-сценариев"],
            ["Spring Boot Test", "Проверка контроллеров и service-слоя backend"],
            ["Vitest + React Testing Library", "Модульное тестирование frontend-компонентов"],
            ["Maven Surefire", "Запуск backend unit-тестов при сборке"],
            ["JaCoCo", "Отчёты покрытия backend-кода"],
            ["GitHub Actions", "Автоматический запуск тестов на push и pull request"],
        ],
    )

    add_heading(doc, "1.3. Применённые техники тест-дизайна", 2)
    add_paragraph(doc, "В модульных тестах были использованы несколько техник тест-дизайна.")
    _number_counter = 0
    add_numbered(doc, "Классы эквивалентности: валидные и невалидные логины, токены, типы предпочтений, наборы полей формы.")
    add_numbered(doc, "Граничные условия: пустые значения, минимальная длина пароля, предельные размеры коллекций и списков.")
    add_numbered(doc, "Негативное тестирование: отсутствие токена, неверные учетные данные, недопустимые комбинации входных данных.")
    add_numbered(doc, "Проверка преобразований данных: корректное маппирование DTO, сохранение и чтение приоритетов и списков.")

    add_heading(doc, "2. Отчёт о прохождении тестов и покрытии кода", 1)
    add_paragraph(
        doc,
        "Модульные тесты встроены в используемые сборщики (`mvn test`, `npm run test:unit`) и автоматически запускаются в "
        "workflow `.github/workflows/project-tests.yml`."
    )
    add_table(
        doc,
        ["Набор", "Количество тестов", "Результат"],
        [
            ["backend unit", "78", "Все тесты пройдены успешно"],
            ["frontend unit", "60", "Все тесты пройдены успешно"],
            ["Итого", "138", "Набор соответствует требованиям автоматизации"],
        ],
    )
    add_paragraph(doc, "Сводка покрытия кода по фактически сформированным отчётам:")
    add_table(
        doc,
        ["Модуль", "Метрика", "Значение"],
        [
            ["backend", "Instruction coverage (JaCoCo)", "91.75%"],
            ["backend", "Line coverage (JaCoCo)", "90.50%"],
            ["backend", "Branch coverage (JaCoCo)", "64.37%"],
            ["frontend", "Lines / Statements (Vitest)", "92.57%"],
            ["frontend", "Branches (Vitest)", "79.02%"],
            ["frontend", "Functions (Vitest)", "57.69%"],
        ],
    )
    add_paragraph(
        doc,
        "Требование по покрытию в 80% выполнено для основной кодовой базы backend и frontend по line/statements coverage. "
        "Отчёты формируются в `backend/target/site/jacoco` и `frontend/coverage`."
    )

    add_heading(doc, "3. Процедура расширения тестового набора", 1)
    add_paragraph(
        doc,
        "Для расширения набора модульных тестов использовался единый сценарий. Ниже приведён пример на добавлении нового "
        "поля предпочтения преподавателя, например `preferredCampus`."
    )
    _number_counter = 0
    add_numbered(doc, "Добавить поле в доменную модель и DTO, а также обновить сериализацию и маппинг в service-слое.")
    add_numbered(doc, "Добавить unit-тесты на getter/setter и корректное преобразование нового поля в backend.")
    add_numbered(doc, "Для frontend добавить unit-тест компонента формы: ввод значения, сохранение и повторное отображение.")
    add_numbered(doc, "Проверить, что новые тесты выполняются теми же командами сборки без дополнительных ручных действий.")
    add_numbered(doc, "Обновить coverage-отчёты и зафиксировать изменение в документации проекта.")

    doc.save(UNIT_OUTPUT)


def build_integration_doc() -> None:
    global _number_counter
    doc = Document(INTEGRATION_TEMPLATE)
    reset_document(doc)
    set_base_styles(doc)
    build_title_page(doc, "Интеграционное и системное тестирование приложения «Teachers Preferences»")

    add_heading(doc, "1. Описание выполненной работы", 1)
    add_paragraph(
        doc,
        "Интеграционное и системное тестирование было выполнено для приложения Teachers Preferences, состоящего из нескольких "
        "взаимодействующих частей: backend, frontend, базы данных PostgreSQL и отдельного модуля schedule-etl. "
        "Основной целью было подтвердить корректную работу пользовательских сценариев на уровне взаимодействия модулей и "
        "на уровне полного пользовательского потока."
    )
    add_table(
        doc,
        ["Компонент", "Назначение", "Роль в тестировании"],
        [
            ["Auth API", "Регистрация и вход пользователя", "Проверяется в интеграционных и системных сценариях"],
            ["Teacher API", "Сохранение и чтение предпочтений преподавателя", "Основной объект интеграционных сценариев"],
            ["Admin API", "Просмотр и экспорт всех предпочтений", "Проверка ролевого доступа и системных сценариев"],
            ["React frontend", "UI преподавателя и администратора", "Проверяется реальными browser E2E тестами"],
            ["PostgreSQL / Docker Compose", "Полноценное runtime-окружение", "Используется в CI для frontend system job"],
        ],
    )

    add_heading(doc, "1.1. Используемые инструменты", 2)
    add_table(
        doc,
        ["Инструмент", "Назначение"],
        [
            ["Spring Boot Test + TestRestTemplate", "Интеграционные и backend system сценарии через реальные HTTP-вызовы"],
            ["Vitest", "Интеграционные frontend-сценарии с контролируемыми mock-зависимостями"],
            ["Playwright", "Реальные browser E2E тесты поверх живого frontend/backend"],
            ["Docker Compose", "Подъём полного окружения в CI для system/E2E"],
            ["GitHub Actions", "Автоматический запуск тестов на pull request, push и вручную"],
            ["Mockito / mocks", "Изоляция внешних зависимостей на интеграционном уровне там, где это оправдано"],
        ],
    )

    add_heading(doc, "2. Тест-план интеграционного тестирования", 1)
    add_paragraph(doc, "Минимальное требование в 10 сценариев выполнено. Основные сценарии интеграционного тестирования:")
    add_table(
        doc,
        ["ID", "Сценарий", "Ожидаемый результат"],
        [
            ["INT-01", "Повторная регистрация пользователя", "Второй запрос отклоняется со статусом 400"],
            ["INT-02", "Вход с неверным паролем", "Система возвращает 401 / отказ во входе"],
            ["INT-03", "Доступ к teacher API без токена", "Система возвращает 403"],
            ["INT-04", "Сохранение semester preference", "Предпочтение сохраняется и возвращается в ответе"],
            ["INT-05", "Сохранение двух записей semester", "При чтении возвращается тот же набор"],
            ["INT-06", "Повторное сохранение session preference", "Старые записи того же типа заменяются новыми"],
            ["INT-07", "Teacher обращается к admin API", "Доступ запрещён, статус 403"],
            ["INT-08", "Изоляция данных semester / session", "Типы не смешиваются между собой"],
            ["INT-09", "Frontend route /teacher/:type", "Композиция маршрутов корректно рендерит форму"],
            ["INT-10", "Frontend unknown route", "Показывается страница 404"],
        ],
    )
    add_paragraph(
        doc,
        "На интеграционном уровне используются заглушки для изоляции окружения только там, где это необходимо: "
        "например, во frontend integration-тестах замещаются API-вызовы и отдельные React-компоненты. "
        "При этом backend integration-сценарии выполняются через реальный HTTP-слой Spring Boot."
    )

    add_heading(doc, "3. Результаты интеграционных тестов на CI", 1)
    add_table(
        doc,
        ["Набор", "Количество сценариев", "CI-команда"],
        [
            ["backend integration", "10", "mvn -B -ntp test \"-Dtest=*IntegrationTest\""],
            ["frontend integration", "17", "npm run test:integration"],
        ],
    )
    add_paragraph(
        doc,
        "Интеграционные тесты запускаются в GitHub Actions по событию изменения кода. "
        "Отчёты формируются стандартными средствами Maven Surefire и Vitest."
    )

    add_heading(doc, "4. Тест-план системного / End-to-End тестирования", 1)
    add_paragraph(
        doc,
        "Для системного уровня были реализованы как backend system HTTP-сценарии, так и реальные browser E2E тесты "
        "frontend через Playwright. Во frontend system job больше не используются моки API: тесты работают по живому UI "
        "после запуска backend, frontend и PostgreSQL через Docker Compose."
    )
    add_table(
        doc,
        ["ID", "Сценарий", "Ожидаемый результат"],
        [
            ["SYS-01", "Регистрация преподавателя", "Пользователь успешно создаётся через UI/API"],
            ["SYS-02", "Логин преподавателя", "Пользователь попадает на teacher dashboard"],
            ["SYS-03", "Сохранение semester preferences", "Данные сохраняются и доступны после чтения"],
            ["SYS-04", "Сохранение session preferences", "Данные сохраняются и восстанавливаются после reload"],
            ["SYS-05", "Недопустимый логин", "Показывается ошибка аутентификации"],
            ["SYS-06", "Доступ к teacher API без токена", "Система возвращает отказ в доступе"],
            ["SYS-07", "Повторное сохранение semester", "Старый набор заменяется новым"],
            ["SYS-08", "Изоляция semester и session", "Данные разных типов не смешиваются"],
            ["SYS-09", "Teacher не имеет доступа к admin API", "Возвращается 403"],
            ["SYS-10", "Администратор видит сохранённые данные и экспортирует Excel", "UI и backend совместно выполняют бизнес-сценарий"],
        ],
    )

    add_heading(doc, "5. Результаты системных тестов и CI", 1)
    add_table(
        doc,
        ["Набор", "Количество сценариев", "CI-команда / job"],
        [
            ["backend system", "10", "mvn -B -ntp test \"-Dtest=*SystemTest,*SystemE2ETest\""],
            ["frontend Playwright E2E", "4 реальных browser-сценария", "frontend-system job в `.github/workflows/project-tests.yml`"],
        ],
    )
    add_paragraph(
        doc,
        "Workflow `frontend-system` устанавливает Playwright, поднимает приложение через Docker Compose, "
        "ожидает готовности frontend и backend, затем выполняет `npm run test:system` и сохраняет `playwright-report` "
        "в артефакты GitHub Actions. Это устраняет прежнюю ситуацию, когда системные тесты фактически ограничивались jsdom-моками."
    )

    add_heading(doc, "6. Процедура расширения тестового набора", 1)
    add_paragraph(
        doc,
        "При добавлении новой функциональной части, например нового типа предпочтения или дополнительного административного действия, "
        "набор тестов расширяется по следующему шаблону:"
    )
    _number_counter = 0
    add_numbered(doc, "Добавить backend integration-сценарий на взаимодействие нового endpoint с уже существующими модулями.")
    add_numbered(doc, "Добавить frontend integration-тест на композицию маршрута или формы с mock API, если требуется изоляция.")
    add_numbered(doc, "Добавить реальный Playwright E2E сценарий через UI для ключевого пользовательского потока.")
    add_numbered(doc, "Обновить CI workflow и тест-план, если новая функциональность требует отдельного шага подготовки окружения.")

    add_heading(doc, "7. Замечание по CodeRabbit и PR-пайплайну", 1)
    add_paragraph(
        doc,
        "Для устранения зависаний CodeRabbit в pull request была добавлена конфигурация `.coderabbit.yaml`: "
        "авто-review разрешён для любых base branches, отключено ожидание GitHub checks внутри CodeRabbit, "
        "а также исключены из анализа большие артефакты (`target`, `coverage`, `node_modules`, zip-файлы). "
        "Это снижает вероятность бесконечного pending-состояния и ускоряет обработку PR."
    )

    doc.save(INTEGRATION_OUTPUT)


if __name__ == "__main__":
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    build_unit_doc()
    build_integration_doc()
    print(UNIT_OUTPUT)
    print(INTEGRATION_OUTPUT)
